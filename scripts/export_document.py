import argparse
import html
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


BULLET_CHAR = "\u2022"
BLUE = colors.HexColor("#1d4ed8")
HARVARD_FONT_EN = "Times New Roman"
HARVARD_DATE_TAB = 1.05
HARVARD_BODY_SIZE = 10.5
HARVARD_NAME_SIZE = 15
HARVARD_CONTACT_SIZE = 10
HARVARD_HEADING_SIZE = 11
HARVARD_MARGIN = 0.6

HARVARD_SECTIONS = {
    "OBJECTIVE",
    "EDUCATION",
    "EXPERIENCE",
    "LEADERSHIP",
    "ADDITIONAL INFORMATION",
    "求职意向",
    "教育背景",
    "工作经历",
    "领导力",
    "其他信息",
}
CHINESE_HEADINGS = {
    "教育背景",
    "工作经历",
    "项目经历",
    "项目经验",
    "技能",
    "核心技能",
    "经验",
    "教育",
    "项目",
    "求职意向",
    "领导力",
    "其他信息",
}


def contains_cjk(text: str) -> bool:
    return any("\u4e00" <= ch <= "\u9fff" for ch in text)


def is_rule_line(line: str) -> bool:
    stripped = line.strip()
    return bool(stripped) and set(stripped) <= {"_"}


def is_harvard_section(line: str) -> bool:
    stripped = line.strip()
    if stripped in HARVARD_SECTIONS:
        return True
    if stripped in CHINESE_HEADINGS:
        return True
    letters = [ch for ch in stripped if ch.isalpha()]
    return bool(letters) and stripped.upper() == stripped and len(stripped) <= 40


def parse_harvard_date_line(line: str):
    patterns = [
        r"^(\d{1,2}/\d{2}-\d{1,2}/\d{2})\s{2,}(.+)$",
        r"^(\d{1,2}/\d{4}-\d{1,2}/\d{4})\s{2,}(.+)$",
        r"^([A-Za-z]{3,9}\s+\d{4}\s*-\s*[A-Za-z]{3,9}\s+\d{4})\s{2,}(.+)$",
        r"^([A-Za-z]{3,9}\s+\d{4}\s*-\s*Present)\s{2,}(.+)$",
    ]
    for pattern in patterns:
        match = re.match(pattern, line)
        if match:
            return match.group(1), match.group(2)
    return None


# Many resumes (including most uploaded ones) place the date at the END of an
# entry line, right-aligned, e.g. "Cornell University, Ithaca, NY    May 2026"
# or "Industrial Engineer Intern, Siemens Energy    June 2024 - August 2024".
# Detect that trailing date so we can keep the title on the left and right-align
# the date, matching the candidate's original layout instead of forcing a new one.
_MONTH = (
    r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
    r"Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\.?"
)
_END_WORD = r"(?:Present|Current|present|current|Now|now)"
# One side of a date: "May 2026", "Nov,2025", "2026", "June", or "Present".
_SIDE = rf"(?:(?:{_MONTH}\s*,?\s*)?(?:\d{{4}}|{_END_WORD})|{_MONTH})"
# A full date: one side, optionally a range to a second side.
_DATE = rf"{_SIDE}(?:\s*(?:[-–—]|to)\s*{_SIDE})?"
_TRAILING_DATE_RE = re.compile(rf"^(?P<text>.*?\S)\s+(?P<date>{_DATE})\s*$")


def parse_harvard_trailing_date(line: str):
    match = _TRAILING_DATE_RE.match(line.strip())
    if not match:
        return None
    date = match.group("date").strip()
    text = match.group("text").strip()
    # Require a real year or an open-ended marker, so we don't mistake a stray
    # month word or a trailing number for a date.
    if not (re.search(r"\d{4}", date) or re.search(_END_WORD, date)):
        return None
    if not text:
        return None
    return text, date


def parse_harvard_resume(content: str):
    lines = [line.rstrip() for line in content.splitlines()]
    non_empty = [line for line in lines if line.strip()]
    if not non_empty:
        return {"name": "", "contact": "", "blocks": []}

    name = non_empty[0].strip()
    contact = ""
    body_start = 1
    if len(non_empty) > 1 and not is_harvard_section(non_empty[1]) and not is_rule_line(non_empty[1]):
        contact = non_empty[1].strip()
        body_start = 2

    blocks = []
    current_section = None
    current_items = []

    def flush_section():
        nonlocal current_section, current_items
        if current_section:
            blocks.append({"type": "section", "title": current_section, "items": current_items})
        current_section = None
        current_items = []

    for raw in non_empty[body_start:]:
        line = raw.rstrip()
        if is_rule_line(line):
            continue
        if is_harvard_section(line):
            flush_section()
            current_section = line.strip()
            continue
        if not current_section:
            continue

        date_match = parse_harvard_date_line(line)
        if date_match:
            current_items.append({"kind": "dated", "date": date_match[0], "text": date_match[1], "bold": True})
            continue

        stripped = line.lstrip()
        indent = len(line) - len(stripped)
        if stripped.startswith((BULLET_CHAR, "-", "*", "\u00b7")):
            current_items.append({"kind": "bullet", "text": clean_bullet(stripped)})
            continue
        if indent >= 4:
            current_items.append({"kind": "indent", "text": stripped, "bold": False})
            continue

        trailing = parse_harvard_trailing_date(stripped)
        if trailing:
            current_items.append(
                {"kind": "entry", "text": trailing[0], "date": trailing[1], "bold": True}
            )
            continue

        current_items.append({"kind": "body", "text": stripped})

    flush_section()
    return {"name": name, "contact": contact, "blocks": blocks}


def add_paragraph_bottom_border(paragraph) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def build_harvard_resume_docx(
    content: str,
    output_path: Path,
    language: str,
    layout: dict[str, float] | None = None,
) -> None:
    parsed = parse_harvard_resume(content)
    layout = layout or harvard_layout_profile(content)
    font_name = HARVARD_FONT_EN if language == "en" else "Times New Roman"
    body_size = layout["body_size"]
    name_size = layout["name_size"]
    contact_size = layout["contact_size"]
    heading_size = layout["heading_size"]
    margin = layout["margin"]
    section_gap = layout["section_gap"]

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(margin)
    section.bottom_margin = Inches(margin)
    section.left_margin = Inches(margin)
    section.right_margin = Inches(margin)

    if parsed["name"]:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(parsed["name"].upper())
        configure_docx_run(run, font_name, name_size, bold=True)

    if parsed["contact"]:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(section_gap + 2)
        run = paragraph.add_run(parsed["contact"])
        configure_docx_run(run, font_name, contact_size)

    for block in parsed["blocks"]:
        header = document.add_paragraph()
        header.paragraph_format.space_before = Pt(section_gap + 2)
        header.paragraph_format.space_after = Pt(1)
        run = header.add_run(block["title"])
        configure_docx_run(run, font_name, heading_size, bold=True)
        add_paragraph_bottom_border(header)

        for item in block["items"]:
            if item["kind"] == "dated":
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(HARVARD_DATE_TAB))
                paragraph.paragraph_format.space_after = Pt(0)
                run_date = paragraph.add_run(item["date"])
                configure_docx_run(run_date, font_name, body_size)
                paragraph.add_run("\t")
                run_text = paragraph.add_run(item["text"])
                configure_docx_run(run_text, font_name, body_size, bold=item.get("bold", False))
            elif item["kind"] == "entry":
                # Title on the left, date right-aligned to the right margin, matching
                # the common uploaded-resume layout.
                paragraph = document.add_paragraph()
                right_tab = 8.5 - 2 * margin
                paragraph.paragraph_format.tab_stops.add_tab_stop(
                    Inches(right_tab), WD_TAB_ALIGNMENT.RIGHT
                )
                paragraph.paragraph_format.space_after = Pt(0)
                run_text = paragraph.add_run(item["text"])
                configure_docx_run(run_text, font_name, body_size, bold=item.get("bold", True))
                paragraph.add_run("\t")
                run_date = paragraph.add_run(item["date"])
                configure_docx_run(run_date, font_name, body_size, bold=item.get("bold", True))
            elif item["kind"] == "indent":
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.left_indent = Inches(HARVARD_DATE_TAB)
                paragraph.paragraph_format.space_after = Pt(0)
                run = paragraph.add_run(item["text"])
                configure_docx_run(run, font_name, body_size, bold=item.get("bold", False))
            elif item["kind"] == "bullet":
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.left_indent = Inches(HARVARD_DATE_TAB)
                paragraph.paragraph_format.first_line_indent = Inches(-0.14)
                paragraph.paragraph_format.space_after = Pt(0)
                run = paragraph.add_run(f"{BULLET_CHAR} {item['text']}")
                configure_docx_run(run, font_name, body_size)
            else:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(1)
                run = paragraph.add_run(item["text"])
                configure_docx_run(run, font_name, body_size)

    document.save(output_path)


def build_harvard_resume_pdf(
    content: str,
    output_path: Path,
    language: str,
    layout: dict[str, float] | None = None,
) -> None:
    parsed = parse_harvard_resume(content)
    layout = layout or harvard_layout_profile(content)
    cjk = contains_cjk(content) or language == "zh"
    base_font = "STSong-Light" if cjk else "Times-Roman"
    bold_font = base_font if cjk else "Times-Bold"
    if cjk:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))

    body_size = layout["body_size"]
    name_size = layout["name_size"]
    contact_size = layout["contact_size"]
    heading_size = layout["heading_size"]
    margin = layout["margin"]
    section_gap = layout["section_gap"]
    line_gap = layout["line_gap"]

    styles = {
        "name": ParagraphStyle(
            "HarvardName",
            fontName=bold_font,
            fontSize=name_size,
            leading=name_size + 2,
            alignment=1,
            textColor=colors.black,
            spaceAfter=4,
        ),
        "contact": ParagraphStyle(
            "HarvardContact",
            fontName=base_font,
            fontSize=contact_size,
            leading=contact_size * line_gap + 1,
            alignment=1,
            textColor=colors.black,
            spaceAfter=section_gap + 2,
        ),
        "heading": ParagraphStyle(
            "HarvardHeading",
            fontName=bold_font,
            fontSize=heading_size,
            leading=heading_size * line_gap + 1,
            textColor=colors.black,
            spaceBefore=section_gap + 2,
            spaceAfter=2,
        ),
        "dated": ParagraphStyle(
            "HarvardDated",
            fontName=bold_font,
            fontSize=body_size,
            leading=body_size * line_gap,
            leftIndent=HARVARD_DATE_TAB * inch,
            firstLineIndent=-HARVARD_DATE_TAB * inch,
            textColor=colors.black,
            spaceAfter=0,
        ),
        "indent": ParagraphStyle(
            "HarvardIndent",
            fontName=base_font,
            fontSize=body_size,
            leading=body_size * line_gap,
            leftIndent=HARVARD_DATE_TAB * inch,
            textColor=colors.black,
            spaceAfter=0,
        ),
        "bullet": ParagraphStyle(
            "HarvardBullet",
            fontName=base_font,
            fontSize=body_size,
            leading=body_size * line_gap,
            leftIndent=HARVARD_DATE_TAB * inch,
            firstLineIndent=-0.14 * inch,
            textColor=colors.black,
            spaceAfter=0,
        ),
        "body": ParagraphStyle(
            "HarvardBody",
            fontName=base_font,
            fontSize=body_size,
            leading=body_size * line_gap,
            textColor=colors.black,
            spaceAfter=1,
        ),
        "entry_title": ParagraphStyle(
            "HarvardEntryTitle",
            fontName=bold_font,
            fontSize=body_size,
            leading=body_size * line_gap,
            textColor=colors.black,
            spaceAfter=0,
        ),
        "entry_date": ParagraphStyle(
            "HarvardEntryDate",
            fontName=bold_font,
            fontSize=body_size,
            leading=body_size * line_gap,
            textColor=colors.black,
            alignment=2,  # right
            spaceAfter=0,
        ),
    }

    content_width = (8.5 - 2 * margin) * inch

    story = []
    if parsed["name"]:
        story.append(Paragraph(html.escape(parsed["name"].upper()), styles["name"]))
    if parsed["contact"]:
        story.append(Paragraph(html.escape(parsed["contact"]), styles["contact"]))

    for block in parsed["blocks"]:
        story.append(Paragraph(html.escape(block["title"]), styles["heading"]))
        for item in block["items"]:
            if item["kind"] == "dated":
                story.append(
                    Paragraph(
                        f"{html.escape(item['date'])}&nbsp;&nbsp;&nbsp;&nbsp;{html.escape(item['text'])}",
                        styles["dated"],
                    )
                )
            elif item["kind"] == "entry":
                # Two-cell row: title left, date right-aligned. Size the table to the
                # frame's INNER width (the document frame insets flowables by 6pt of
                # padding) and left-align it, so the title's left edge lines up exactly
                # with the section headings and body paragraphs, and the date's right
                # edge lines up with the wrapped body text.
                row_width = content_width - 12
                date_w = pdfmetrics.stringWidth(item["date"], bold_font, body_size) + 4
                date_w = min(date_w, row_width * 0.5)
                title_cell = Paragraph(html.escape(item["text"]), styles["entry_title"])
                date_cell = Paragraph(html.escape(item["date"]), styles["entry_date"])
                table = Table(
                    [[title_cell, date_cell]],
                    colWidths=[row_width - date_w, date_w],
                )
                table.hAlign = "LEFT"
                table.setStyle(
                    TableStyle(
                        [
                            ("LEFTPADDING", (0, 0), (-1, -1), 0),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                            ("TOPPADDING", (0, 0), (-1, -1), 0),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ]
                    )
                )
                story.append(table)
            elif item["kind"] == "bullet":
                story.append(Paragraph(f"{BULLET_CHAR} {html.escape(item['text'])}", styles["bullet"]))
            elif item["kind"] == "indent":
                story.append(Paragraph(html.escape(item["text"]), styles["indent"]))
            else:
                story.append(Paragraph(html.escape(item["text"]), styles["body"]))

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        leftMargin=margin * inch,
        rightMargin=margin * inch,
        topMargin=margin * inch,
        bottomMargin=margin * inch,
    )
    doc.build(story)


def is_heading(line: str) -> bool:
    return is_harvard_section(line)


def is_bullet(line: str) -> bool:
    stripped = line.lstrip()
    return stripped.startswith((BULLET_CHAR, "-", "*", "\u00b7"))


def clean_bullet(line: str) -> str:
    stripped = line.lstrip()
    while stripped.startswith((BULLET_CHAR, "-", "*", "\u00b7")):
        stripped = stripped[1:].lstrip()
    return stripped


def looks_like_entry(line: str) -> bool:
    stripped = line.strip()
    if not stripped or is_heading(stripped) or is_bullet(stripped):
        return False
    if " | " in stripped:
        return True
    if re.search(r"\b(19|20)\d{2}\b", stripped):
        return True
    if re.search(r"\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\b", stripped, re.I):
        return True
    return len(stripped) <= 90 and stripped.count(",") <= 2


def harvard_layout_profile(
    content: str,
    *,
    comfortable: bool = False,
    force_tight: bool = False,
    ultra_tight: bool = False,
) -> dict[str, float]:
    non_empty_lines = [line for line in content.splitlines() if line.strip()]
    char_count = len(content)
    line_count = len(non_empty_lines)

    if ultra_tight or line_count > 50 or char_count > 3600:
        return {
            "name_size": 14.0,
            "contact_size": 8.4,
            "body_size": 8.8,
            "heading_size": 9.6,
            "margin": 0.38,
            "line_gap": 1.0,
            "section_gap": 2,
        }
    if force_tight or line_count > 44 or char_count > 3000:
        return {
            "name_size": 14.5,
            "contact_size": 8.6,
            "body_size": 9.2,
            "heading_size": 10.0,
            "margin": 0.44,
            "line_gap": 1.02,
            "section_gap": 2,
        }
    if line_count > 38 or char_count > 2500:
        return {
            "name_size": 15.0,
            "contact_size": 8.8,
            "body_size": 9.8,
            "heading_size": 10.2,
            "margin": 0.5,
            "line_gap": 1.03,
            "section_gap": 3,
        }
    if comfortable or line_count < 24:
        return {
            "name_size": 17.2,
            "contact_size": 10.6,
            "body_size": 12.0,
            "heading_size": 11.8,
            "margin": 0.74,
            "line_gap": 1.16,
            "section_gap": 6,
        }
    if line_count < 30:
        return {
            "name_size": 16.6,
            "contact_size": 10.2,
            "body_size": 11.4,
            "heading_size": 11.3,
            "margin": 0.69,
            "line_gap": 1.13,
            "section_gap": 5,
        }
    if line_count < 36:
        return {
            "name_size": 16.0,
            "contact_size": 9.8,
            "body_size": 10.8,
            "heading_size": 10.9,
            "margin": 0.64,
            "line_gap": 1.1,
            "section_gap": 5,
        }
    return {
        "name_size": HARVARD_NAME_SIZE,
        "contact_size": HARVARD_CONTACT_SIZE,
        "body_size": HARVARD_BODY_SIZE,
        "heading_size": HARVARD_HEADING_SIZE,
        "margin": HARVARD_MARGIN,
        "line_gap": 1.06,
        "section_gap": 4,
    }


def compact_resume_profile(content: str) -> dict[str, float]:
    return harvard_layout_profile(content)


def harvard_layout_for_scale(scale: float) -> dict[str, float]:
    """A continuously scalable resume layout. Larger scale => larger font, line
    spacing, margins, and section gaps, which consumes more vertical space. We
    search for the largest scale that still fits one page, so the resume fills
    the page instead of leaving the bottom empty."""
    if scale > 1.0:
        margin = min(0.95, 0.55 + 0.45 * (scale - 1.0))
    else:
        margin = max(0.45, 0.55 + 0.30 * (scale - 1.0))
    return {
        "name_size": round(13.5 * scale, 1),
        "contact_size": round(9.0 * scale, 1),
        "body_size": round(10.2 * scale, 1),
        "heading_size": round(10.6 * scale, 1),
        "margin": round(margin, 3),
        "line_gap": round(max(1.0, min(1.34, 1.02 + 0.40 * (scale - 1.0))), 3),
        "section_gap": max(2, round(4 * scale)),
    }


def best_single_page_layout(
    content: str,
    language: str,
    *,
    min_scale: float = 0.80,
    max_scale: float = 1.60,
) -> dict[str, float]:
    """Binary-search the largest scale whose resume still fits on one page.
    Page count grows monotonically with scale, so binary search is safe."""
    import tempfile

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as handle:
        probe_path = Path(handle.name)

    def fits(scale: float) -> bool:
        try:
            build_harvard_resume_pdf(
                content, probe_path, language, layout=harvard_layout_for_scale(scale)
            )
            return resume_pdf_page_count(probe_path) <= 1
        except Exception:
            return False

    try:
        # If even the smallest scale overflows, just use it (best effort).
        if not fits(min_scale):
            return harvard_layout_for_scale(min_scale)
        lo, hi = min_scale, max_scale
        # 7 iterations narrows the window to < 0.01.
        for _ in range(7):
            mid = (lo + hi) / 2
            if fits(mid):
                lo = mid
            else:
                hi = mid
        return harvard_layout_for_scale(round(lo, 3))
    finally:
        try:
            probe_path.unlink()
        except OSError:
            pass


def split_name_and_body(content: str) -> tuple[str, str, list[str]]:
    raw_lines = [line.rstrip() for line in content.splitlines()]
    name = ""
    contact = ""
    body: list[str] = []
    seen_name = False
    seen_contact = False

    for line in raw_lines:
        stripped = line.strip()
        if stripped and not seen_name:
            name = stripped
            seen_name = True
            continue
        if stripped and not seen_contact and not is_heading(stripped):
            contact = stripped
            seen_contact = True
            continue
        body.append(stripped)

    return name, contact, body


def configure_docx_run(run, font_name: str, size: float, bold: bool = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)


def build_resume_docx(content: str, output_path: Path, language: str) -> None:
    # Use the same fill-the-page layout the PDF uses so the Word version is also
    # sized to fill roughly a full page rather than leaving the bottom empty.
    layout = best_single_page_layout(content, language)
    build_harvard_resume_docx(content, output_path, language, layout=layout)


def build_resume_pdf(
    content: str,
    output_path: Path,
    language: str,
    layout: dict[str, float] | None = None,
) -> None:
    build_harvard_resume_pdf(content, output_path, language, layout=layout)


def resume_pdf_page_count(output_path: Path) -> int:
    return len(PdfReader(str(output_path)).pages)


def build_single_page_resume_pdf(content: str, output_path: Path, language: str) -> None:
    layout = best_single_page_layout(content, language)
    build_resume_pdf(content, output_path, language, layout=layout)


def parse_cover_letter(content: str):
    lines = [line.rstrip() for line in content.splitlines()]
    greeting_index = next(
        (index for index, line in enumerate(lines) if line.strip().lower().startswith("dear ")),
        None,
    )

    if greeting_index is None:
        paragraphs = [part.strip() for part in re.split(r"\n\s*\n", content) if part.strip()]
        return {
            "header_lines": [],
            "header_groups": [],
            "body_paragraphs": paragraphs,
            "closing_lines": [],
        }

    closing_start = None
    for index in range(len(lines) - 1, greeting_index, -1):
        stripped = lines[index].strip()
        if stripped.lower() in {"best regard,", "best regards,", "sincerely,", "thank you,"}:
            closing_start = index
            break

    # Preserve blank-line boundaries in the header so the sender block, the date,
    # and the recipient block stay visually separated. header_groups is a list of
    # blocks (each a list of lines); header_lines is the flattened version.
    header_groups = []
    current_group = []
    for line in lines[:greeting_index]:
        if line.strip():
            current_group.append(line.strip())
        elif current_group:
            header_groups.append(current_group)
            current_group = []
    if current_group:
        header_groups.append(current_group)
    header_lines = [line for group in header_groups for line in group]

    if closing_start is None:
        body_lines = lines[greeting_index:]
        closing_lines = []
    else:
        body_lines = lines[greeting_index:closing_start]
        closing_lines = [line.strip() for line in lines[closing_start:] if line.strip()]

    body_paragraphs = []
    current = []
    for line in body_lines:
        stripped = line.strip()
        if stripped:
            current.append(stripped)
        elif current:
            body_paragraphs.append(" ".join(current))
            current = []
    if current:
        body_paragraphs.append(" ".join(current))

    return {
        "header_lines": header_lines,
        "header_groups": header_groups,
        "body_paragraphs": body_paragraphs,
        "closing_lines": closing_lines,
    }


def build_cover_docx(content: str, output_path: Path, language: str) -> None:
    font_name = "Arial" if language == "en" else "PingFang SC"
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    style = document.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    parsed = parse_cover_letter(content)

    header_groups = parsed["header_groups"]
    for group_index, group in enumerate(header_groups):
        last_group = group_index == len(header_groups) - 1
        for line_index, line in enumerate(group):
            last_in_group = line_index == len(group) - 1
            paragraph = document.add_paragraph()
            # Add a visible gap after each block (e.g. between the sender info and
            # the recipient info), but keep lines within a block tight.
            gap = 8 if (last_in_group and not last_group) else (3 if line.startswith("RE:") else 1)
            paragraph.paragraph_format.space_after = Pt(gap)
            paragraph.paragraph_format.line_spacing = 1.1
            run = paragraph.add_run(line)
            configure_docx_run(run, font_name, 11, bold=line.startswith("RE:"))

    if parsed["header_lines"]:
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)

    for paragraph_text in parsed["body_paragraphs"]:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(10)
        paragraph.paragraph_format.line_spacing = 1.18
        run = paragraph.add_run(paragraph_text)
        configure_docx_run(run, font_name, 11)

    if parsed["closing_lines"]:
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)
        for line in parsed["closing_lines"]:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.1
            run = paragraph.add_run(line)
            configure_docx_run(run, font_name, 11)

    document.save(output_path)


def build_cover_pdf(content: str, output_path: Path, language: str) -> None:
    cjk = contains_cjk(content) or language == "zh"
    base_font = "STSong-Light" if cjk else "Helvetica"
    if cjk:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))

    body_style = ParagraphStyle(
        "CoverBody",
        fontName=base_font,
        fontSize=11,
        leading=15,
        textColor=colors.HexColor("#0f172a"),
        spaceAfter=10,
    )
    header_style = ParagraphStyle(
        "CoverHeader",
        fontName=base_font,
        fontSize=11,
        leading=13,
        textColor=colors.HexColor("#0f172a"),
        spaceAfter=2,
    )
    header_bold_style = ParagraphStyle(
        "CoverHeaderBold",
        parent=header_style,
        fontName="STSong-Light" if cjk else "Helvetica-Bold",
        spaceAfter=3,
    )

    parsed = parse_cover_letter(content)
    story = []
    header_groups = parsed["header_groups"]
    for group_index, group in enumerate(header_groups):
        for line in group:
            style = header_bold_style if line.startswith("RE:") else header_style
            story.append(Paragraph(html.escape(line), style))
        # Visible gap between blocks (e.g. sender info vs recipient info).
        if group_index < len(header_groups) - 1:
            story.append(Spacer(1, 7))
    if parsed["header_lines"]:
        story.append(Spacer(1, 8))

    for block in parsed["body_paragraphs"]:
        story.append(Paragraph(html.escape(block), body_style))
        story.append(Spacer(1, 2))

    if parsed["closing_lines"]:
        story.append(Spacer(1, 6))
        for line in parsed["closing_lines"]:
            story.append(Paragraph(html.escape(line), header_style))

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        leftMargin=0.9 * inch,
        rightMargin=0.9 * inch,
        topMargin=0.85 * inch,
        bottomMargin=0.85 * inch,
    )
    doc.build(story)


def build_pdf(content: str, output_path: Path, document_type: str, language: str) -> None:
    if document_type == "resume":
        build_single_page_resume_pdf(content, output_path, language)
        return
    build_cover_pdf(content, output_path, language)


def build_docx(content: str, output_path: Path, document_type: str, language: str) -> None:
    if document_type == "resume":
        build_resume_docx(content, output_path, language)
        return
    build_cover_docx(content, output_path, language)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    payload = json.loads(Path(args.input).read_text(encoding="utf-8"))
    content = payload["content"].strip()
    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if payload["format"] == "docx":
        build_docx(content, output_path, payload["document_type"], payload["language"])
    elif payload["format"] == "pdf":
        build_pdf(content, output_path, payload["document_type"], payload["language"])
    else:
        raise ValueError("Unsupported export format")


if __name__ == "__main__":
    main()
